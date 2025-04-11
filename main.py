import os
import io
import requests
from bs4 import BeautifulSoup
import json
from moviepy import ImageSequenceClip, AudioFileClip, concatenate_videoclips
import sys
import re
from PIL import Image, ImageDraw, ImageFont
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from tkinter import font as tkfont
import sv_ttk
from datetime import datetime
import logging
import subprocess
import numpy as np
import moviepy.video.fx as vfx
from docx import Document
import time
import pickle
from pathlib import Path

try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    logging.warning("google.generativeai not available")

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    filename='ai_content_creator.log'
)

class VideoCreatorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("AI Content Creation Tool - CATSMOKER PRO")
        self.root.geometry("500x700")
        self.root.resizable(False, False)
        sv_ttk.set_theme("dark")
        self.title_font = tkfont.Font(family="Segoe UI", size=24, weight="bold")
        self.subtitle_font = tkfont.Font(family="Segoe UI", size=12)
        self.button_font = tkfont.Font(family="Segoe UI", size=12, weight="bold")
        self.api_keys_file = "api_keys.pkl"
        self.saved_api_keys = self.load_api_keys()
        self.setup_ui()
        self.create_output_folder()
        self.running = False
    
    def load_api_keys(self):
        try:
            if os.path.exists(self.api_keys_file):
                with open(self.api_keys_file, 'rb') as f:
                    return pickle.load(f)
        except Exception as e:
            logging.error(f"Error loading API keys: {str(e)}")
        return {'gemini': '', 'elevenlabs': ''}
    
    def save_api_keys(self):
        try:
            with open(self.api_keys_file, 'wb') as f:
                pickle.dump({
                    'gemini': self.gemini_entry.get().strip(),
                    'elevenlabs': self.eleven_entry.get().strip()
                }, f)
        except Exception as e:
            logging.error(f"Error saving API keys: {str(e)}")

    def create_output_folder(self):
        self.output_dir = os.path.join(os.path.expanduser("~"), "Desktop", "AI_Videos_Pro")
        os.makedirs(self.output_dir, exist_ok=True)
    
    def setup_ui(self):
        main_frame = ttk.Frame(self.root, padding=(30, 20))
        main_frame.pack(fill="both", expand=True)
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill="x", pady=(0, 30))
        ttk.Label(
            header_frame,
            text="AI Content Creation Tool Pro",
            font=self.title_font,
            anchor="center"
        ).pack(fill="x")
        ttk.Label(
            header_frame,
            text="Create professional videos with AI",
            font=self.subtitle_font,
            anchor="center"
        ).pack(fill="x")
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.pack(fill="both", expand=True)
        self.setup_api_tab()
        self.setup_content_tab()
        self.setup_effects_tab()
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill="x", pady=20)
        ttk.Button(
            button_frame,
            text="Check API Connection",
            style="Accent.TButton",
            command=self.test_api_connection
        ).pack(side="left", expand=True, padx=10)
        ttk.Button(
            button_frame,
            text="Preview",
            command=self.generate_preview
        ).pack(side="left", expand=True, padx=10)
        ttk.Button(
            button_frame,
            text="Create Content",
            style="Accent.TButton",
            command=self.start_creation_process
        ).pack(side="right", expand=True, padx=10)
        self.status_frame = ttk.Frame(self.root, height=25)
        self.status_frame.pack(fill="x", side="bottom")
        self.status_label = ttk.Label(self.status_frame, text="Ready", anchor="w")
        self.status_label.pack(fill="x", padx=10)
    
    def setup_api_tab(self):
        api_tab = ttk.Frame(self.notebook, padding=20)
        self.notebook.add(api_tab, text="API Settings")
        api_card = ttk.LabelFrame(api_tab, text="API Credentials", padding=20)
        api_card.pack(fill="x", pady=10)
        gemini_frame = ttk.Frame(api_card)
        gemini_frame.pack(fill="x", pady=5)
        ttk.Label(gemini_frame, text="Gemini API Key:", font=self.button_font).pack(side="left", padx=5)
        self.gemini_entry = ttk.Entry(gemini_frame, width=50)
        self.gemini_entry.insert(0, self.saved_api_keys['gemini'])
        self.gemini_entry.pack(side="right", expand=True, fill="x", padx=5)
        eleven_frame = ttk.Frame(api_card)
        eleven_frame.pack(fill="x", pady=5)
        ttk.Label(eleven_frame, text="ElevenLabs API Key:", font=self.button_font).pack(side="left", padx=5)
        self.eleven_entry = ttk.Entry(eleven_frame, width=50)
        self.eleven_entry.insert(0, self.saved_api_keys['elevenlabs'])
        self.eleven_entry.pack(side="right", expand=True, fill="x", padx=5)
    
    def setup_content_tab(self):
        content_tab = ttk.Frame(self.notebook, padding=20)
        self.notebook.add(content_tab, text="Content Settings")
        content_card = ttk.LabelFrame(content_tab, text="Content Configuration", padding=20)
        content_card.pack(fill="x", pady=10)
        content_type_frame = ttk.Frame(content_card)
        content_type_frame.pack(fill="x", pady=10)
        ttk.Label(content_type_frame, text="Content Type:", font=self.button_font).pack(side="left", padx=5)
        self.content_types = ["Sports", "Technology", "Health & Fitness", "Business", "Travel", "Education", "Entertainment", "News", "Science", "Motivational"]
        self.content_type_combo = ttk.Combobox(content_type_frame, values=self.content_types, font=self.subtitle_font)
        self.content_type_combo.pack(side="right", expand=True, fill="x", padx=5)
        self.content_type_combo.current(0)
        style_frame = ttk.Frame(content_card)
        style_frame.pack(fill="x", pady=10)
        ttk.Label(style_frame, text="Content Style:", font=self.button_font).pack(side="left", padx=5)
        self.style_options = ["Professional", "Casual", "Educational", "Entertaining", "Inspirational"]
        self.style_combo = ttk.Combobox(style_frame, values=self.style_options, font=self.subtitle_font)
        self.style_combo.pack(side="right", expand=True, fill="x", padx=5)
        self.style_combo.current(0)
        duration_frame = ttk.Frame(content_card)
        duration_frame.pack(fill="x", pady=10)
        ttk.Label(duration_frame, text="Video Duration:", font=self.button_font).pack(side="left", padx=5)
        self.duration_options = ["30 seconds", "60 seconds", "90 seconds", "2 minutes", "3 minutes", "5 minutes"]
        self.duration_combo = ttk.Combobox(duration_frame, values=self.duration_options, font=self.subtitle_font)
        self.duration_combo.pack(side="right", expand=True, fill="x", padx=5)
        self.duration_combo.current(1)
        voice_frame = ttk.Frame(content_card)
        voice_frame.pack(fill="x", pady=10)
        ttk.Label(voice_frame, text="Voice Type:", font=self.button_font).pack(side="left", padx=5)
        self.voice_options = ["Professional Male", "Inspirational Female", "Young Male", "Young Female", "Narrator"]
        self.voice_combo = ttk.Combobox(voice_frame, values=self.voice_options, font=self.subtitle_font)
        self.voice_combo.pack(side="right", expand=True, fill="x", padx=5)
        self.voice_combo.current(0)
        output_frame = ttk.Frame(content_card)
        output_frame.pack(fill="x", pady=10)
        ttk.Label(output_frame, text="Save Location:", font=self.button_font).pack(side="left", padx=5)
        self.output_var = tk.StringVar()
        self.output_var.set(os.path.join(os.path.expanduser("~"), "Desktop", "AI_Videos_Pro"))
        ttk.Entry(output_frame, textvariable=self.output_var, font=self.subtitle_font).pack(side="left", expand=True, fill="x", padx=5)
        ttk.Button(
            output_frame,
            text="Browse...",
            command=self.select_output_folder
        ).pack(side="right", padx=5)
    
    def setup_effects_tab(self):
        effects_tab = ttk.Frame(self.notebook, padding=20)
        self.notebook.add(effects_tab, text="Video Effects")
        effects_card = ttk.LabelFrame(effects_tab, text="Video Effects Settings", padding=20)
        effects_card.pack(fill="x", pady=10)
        zoom_frame = ttk.Frame(effects_card)
        zoom_frame.pack(fill="x", pady=5)
        ttk.Label(zoom_frame, text="Zoom Effect:", font=self.button_font).pack(side="left", padx=5)
        self.zoom_var = tk.DoubleVar(value=1.03)
        ttk.Scale(zoom_frame, from_=1.0, to=1.2, variable=self.zoom_var, orient="horizontal").pack(side="right", expand=True, fill="x", padx=5)
        transition_frame = ttk.Frame(effects_card)
        transition_frame.pack(fill="x", pady=5)
        ttk.Label(transition_frame, text="Transition:", font=self.button_font).pack(side="left", padx=5)
        self.transition_options = ["Crossfade", "Slide", "Fade to Black", "None"]
        self.transition_combo = ttk.Combobox(transition_frame, values=self.transition_options)
        self.transition_combo.current(0)
        self.transition_combo.pack(side="right", expand=True, fill="x", padx=5)
        duration_frame = ttk.Frame(effects_card)
        duration_frame.pack(fill="x", pady=5)
        ttk.Label(duration_frame, text="Image Duration (sec):", font=self.button_font).pack(side="left", padx=5)
        self.img_duration_var = tk.IntVar(value=5)
        ttk.Entry(duration_frame, textvariable=self.img_duration_var, width=5).pack(side="right", padx=5)
    
    def select_output_folder(self):
        folder = filedialog.askdirectory(title="Select Output Folder")
        if folder:
            self.output_var.set(folder)
            self.output_dir = folder
    
    def test_api_connection(self):
        try:
            if not GEMINI_AVAILABLE:
                raise ImportError("google.generativeai package not installed")
            genai.configure(api_key=self.gemini_entry.get().strip())
            models = genai.list_models()
            if not models:
                raise ValueError("No models found - check your API key")
            headers = {"xi-api-key": self.eleven_entry.get().strip()}
            response = requests.get("https://api.elevenlabs.io/v1/user", headers=headers, timeout=10)
            response.raise_for_status()
            self.save_api_keys()
            messagebox.showinfo("Success", "Both API connections are working!")
        except Exception as e:
            self.show_error(f"API Connection Failed:\n{str(e)}")
    
    def generate_preview(self):
        if self.running:
            return
        try:
            self.running = True
            if not self.validate_inputs():
                self.running = False
                return
            self.progress_window = tk.Toplevel(self.root)
            self.progress_window.title("Creating Preview")
            self.progress_window.geometry("400x200")
            self.progress_window.protocol("WM_DELETE_WINDOW", self.cancel_creation)
            frame = ttk.Frame(self.progress_window, padding=20)
            frame.pack(fill="both", expand=True)
            self.progress_label = ttk.Label(frame, text="Creating 10-second preview...", font=self.subtitle_font)
            self.progress_label.pack(pady=5)
            self.progress_bar = ttk.Progressbar(frame, orient="horizontal", length=300, mode="determinate")
            self.progress_bar.pack(pady=10)
            self.cancel_button = ttk.Button(
                frame,
                text="Cancel",
                command=self.cancel_creation
            )
            self.cancel_button.pack(pady=5)
            self.root.after(100, self._generate_preview_content)
        except Exception as e:
            self.show_error(f"Failed to start preview: {str(e)}")
            self.running = False
    
    def _generate_preview_content(self):
        try:
            self.update_progress(20, "Generating preview script...")
            script = self.generate_script()[:500]
            if not script or not self.running:
                return
            self.update_progress(40, "Generating preview voiceover...")
            voiceover_path = self.generate_voiceover(script[:300])
            if not voiceover_path or not self.running:
                return
            self.update_progress(60, "Downloading preview images...")
            image_folder = "temp_preview_images"
            os.makedirs(image_folder, exist_ok=True)
            self.download_images(script, image_folder, max_images=2)
            if not self.running:
                return
            self.update_progress(80, "Creating preview video...")
            preview_path = os.path.join(self.output_dir, "preview.mp4")
            audio_clip = AudioFileClip(voiceover_path)
            target_duration = min(10, audio_clip.duration)
            image_files = [os.path.join(image_folder, f) for f in os.listdir(image_folder) 
                          if f.lower().endswith(('.png', '.jpg', '.jpeg'))]
            if image_files:
                clips = []
                duration_per_image = target_duration / len(image_files)
                for img in image_files:
                    clip = ImageSequenceClip([img], durations=[duration_per_image])
                    clip = clip.resize(height=1080)
                    clips.append(clip)
                final_clip = concatenate_videoclips(clips)
                final_clip = final_clip.set_audio(audio_clip.subclip(0, target_duration))
                final_clip.write_videofile(
                    preview_path,
                    fps=24,
                    codec="libx264",
                    audio_codec="aac",
                    threads=4,
                    logger=None
                )
            self.cleanup_temp_files(image_folder, voiceover_path)
            self.update_progress(100, "Preview created!")
            self.show_success(f"Preview created successfully!\n\nSaved to:\n{preview_path}")
            if os.path.exists(preview_path):
                subprocess.Popen(f'explorer "{os.path.dirname(preview_path)}"')
        except Exception as e:
            self.show_error(f"Preview creation failed: {str(e)}")
            logging.exception("Preview creation error")
        finally:
            self.running = False
            if hasattr(self, 'progress_window') and self.progress_window:
                self.progress_window.destroy()
    
    def start_creation_process(self):
        if self.running:
            return
        try:
            self.running = True
            if not self.validate_inputs():
                self.running = False
                return
            self.progress_window = tk.Toplevel(self.root)
            self.progress_window.title("Creating Content")
            self.progress_window.geometry("400x200")
            self.progress_window.protocol("WM_DELETE_WINDOW", self.cancel_creation)
            frame = ttk.Frame(self.progress_window, padding=20)
            frame.pack(fill="both", expand=True)
            self.progress_label = ttk.Label(frame, text="Initializing...", font=self.subtitle_font)
            self.progress_label.pack(pady=5)
            self.progress_bar = ttk.Progressbar(frame, orient="horizontal", length=300, mode="determinate")
            self.progress_bar.pack(pady=10)
            self.cancel_button = ttk.Button(
                frame,
                text="Cancel",
                command=self.cancel_creation
            )
            self.cancel_button.pack(pady=5)
            self.root.after(100, self.create_content)
        except Exception as e:
            self.show_error(f"Failed to start creation process: {str(e)}")
            self.running = False
    
    def validate_inputs(self):
        if not self.gemini_entry.get().strip():
            self.show_error("Please enter your Gemini API key")
            return False
        if not self.eleven_entry.get().strip():
            self.show_error("Please enter your ElevenLabs API key")
            return False
        if not GEMINI_AVAILABLE:
            self.show_error("google.generativeai package is not installed")
            return False
        return True
    
    def cancel_creation(self):
        self.running = False
        if hasattr(self, 'progress_window') and self.progress_window:
            self.progress_window.destroy()
        self.update_status("Creation cancelled")
    
    def update_progress(self, value, message):
        if hasattr(self, 'progress_window') and self.progress_window:
            self.progress_bar['value'] = value
            self.progress_label['text'] = message
            self.progress_window.update()
    
    def update_status(self, message):
        self.status_label['text'] = message
        self.root.update()
    
    def show_error(self, message):
        messagebox.showerror("Error", message)
        logging.error(message)
        self.update_status(f"Error: {message}")
    
    def show_success(self, message):
        messagebox.showinfo("Success", message)
        logging.info(message)
        self.update_status(message)
    
    def create_content(self):
        if not self.running:
            return
        try:
            self.update_progress(10, "Generating script...")
            script = self.generate_script()
            if not script or not self.running:
                return
            if not self.validate_script(script):
                raise ValueError("Generated script failed quality checks")
            self.update_progress(20, "Saving script...")
            script_path = self.save_script_to_docx(script)
            if not script_path or not self.running:
                return
            self.update_progress(30, "Generating voiceover...")
            voiceover_path = self.generate_voiceover(script)
            if not voiceover_path or not self.running:
                return
            self.update_progress(50, "Downloading images...")
            image_folder = "temp_images"
            os.makedirs(image_folder, exist_ok=True)
            self.download_images(script, image_folder)
            if not self.running:
                return
            self.update_progress(80, "Creating video...")
            video_path = self.create_video_with_effects(image_folder, voiceover_path)
            self.cleanup_temp_files(image_folder, voiceover_path)
            self.update_progress(100, "Process completed!")
            self.show_success(f"Video created successfully!\n\nSaved to:\n{video_path}")
            if os.path.exists(video_path):
                subprocess.Popen(f'explorer "{os.path.dirname(video_path)}"')
        except Exception as e:
            self.show_error(f"Content creation failed: {str(e)}")
            logging.exception("Content creation error")
            if messagebox.askretrycancel("Error", "Would you like to try again?"):
                self.start_creation_process()
        finally:
            self.running = False
            if hasattr(self, 'progress_window') and self.progress_window:
                self.progress_window.destroy()
    
    def generate_script(self):
        try:
            api_key = self.gemini_entry.get().strip()
            if not api_key:
                raise ValueError("Gemini API key is required")
            genai.configure(api_key=api_key)
            content_type = self.content_type_combo.get()
            duration = self.duration_combo.get()
            style = self.style_combo.get()
            prompt = f"""Create a professional {duration} YouTube script about {content_type} with this structure:
            1. Engaging hook (first 5-10 seconds)
            2. 3-5 key points with supporting facts
            3. Clear transitions between sections
            4. Call-to-action at the end
            Style: {style}
            Tone: Professional but engaging
            Target audience: General YouTube viewers
            Make it factual, well-structured, and suitable for voiceover narration.
            Avoid filler content and maintain consistent quality throughout."""
            model = genai.GenerativeModel('gemini-1.5-pro')
            response = model.generate_content(
                prompt,
                generation_config={
                    "temperature": 0.3,
                    "top_p": 0.7,
                    "max_output_tokens": 2000
                }
            )
            if not response.text:
                raise ValueError("Empty response from Gemini API")
            return response.text
        except Exception as e:
            error_msg = f"Script generation failed: {str(e)}"
            if "404" in str(e):
                error_msg += "\n\nPossible solutions:\n"
                error_msg += "1. Check your API key is valid\n"
                error_msg += "2. Verify you're using the correct model name\n"
                error_msg += "3. Ensure your account has access to this model"
            self.show_error(error_msg)
            logging.error(f"Script generation error: {str(e)}")
            return None
    
    def validate_script(self, script):
        word_count = len(script.split())
        min_words = 100 if "30 seconds" in self.duration_combo.get() else 150
        if word_count < min_words:
            self.show_error(f"Script too short ({word_count} words). Needs at least {min_words} words.")
            return False
        nonsense_words = ["lorem", "ipsum", "undefined", "example", "placeholder"]
        if any(word in script.lower() for word in nonsense_words):
            self.show_error("Script contains placeholder/nonsense text")
            return False
        return True
    
    def save_script_to_docx(self, script):
        try:
            doc = Document()
            doc.add_heading('AI Generated Script', 0)
            doc.add_paragraph(f"Content Type: {self.content_type_combo.get()}")
            doc.add_paragraph(f"Style: {self.style_combo.get()}")
            doc.add_paragraph(f"Duration: {self.duration_combo.get()}")
            doc.add_paragraph(f"Created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("\n")
            for paragraph in script.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            script_path = os.path.join(self.output_dir, "generated_script.docx")
            doc.save(script_path)
            return script_path
        except Exception as e:
            self.show_error(f"Failed to save script: {str(e)}")
            return None
    
    def generate_voiceover(self, script):
        max_retries = 3
        retry_delay = 5
        timeout_duration = 60
        try:
            api_key = self.eleven_entry.get().strip()
            if not api_key:
                raise ValueError("ElevenLabs API key is required")
            headers = {
                "Accept": "audio/mpeg",
                "Content-Type": "application/json",
                "xi-api-key": api_key
            }
            voice_type = self.voice_combo.get()
            voice_ids = {
                "Professional Male": "pNInz6obpgDQGcFmaJgB",
                "Inspirational Female": "IKne3meq5aSn9XLyUdCD",
                "Young Male": "g5CIjZEefAph4nQFvHAz",
                "Young Female": "jBpfuIE2acCO8z3wKNLl",
                "Narrator": "wViXBPUzp2ZZixB1xQuM"
            }
            voice_id = voice_ids.get(voice_type, "pNInz6obpgDQGcFmaJgB")
            data = {
                "text": script[:5000],
                "model_id": "eleven_multilingual_v2",
                "voice_settings": {
                    "stability": 0.5,
                    "similarity_boost": 0.75,
                    "style": 0.3,
                    "speaker_boost": True
                }
            }
            voiceover_path = os.path.join("temp_voiceover.mp3")
            for attempt in range(max_retries):
                try:
                    self.update_status(f"Generating voiceover (Attempt {attempt + 1}/{max_retries})")
                    response = requests.post(
                        f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}",
                        json=data,
                        headers=headers,
                        timeout=timeout_duration
                    )
                    response.raise_for_status()
                    with open(voiceover_path, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=1024):
                            if chunk and self.running:
                                f.write(chunk)
                            else:
                                raise Exception("Process cancelled")
                    if os.path.exists(voiceover_path) and os.path.getsize(voiceover_path) > 0:
                        self.save_api_keys()
                        return voiceover_path
                    else:
                        raise ValueError("Voiceover file was not created properly")
                except requests.exceptions.Timeout:
                    if attempt < max_retries - 1:
                        time.sleep(retry_delay)
                        continue
                    raise
                except requests.exceptions.RequestException as e:
                    if attempt < max_retries - 1:
                        time.sleep(retry_delay)
                        continue
                    raise
        except Exception as e:
            error_msg = f"Voiceover generation failed after {max_retries} attempts"
            if isinstance(e, requests.exceptions.Timeout):
                error_msg += "\n\nThe server took too long to respond."
                error_msg += "\nPossible solutions:"
                error_msg += "\n1. Check your internet connection"
                error_msg += "\n2. Try again later (server might be busy)"
                error_msg += "\n3. Use shorter text for voiceover"
            elif isinstance(e, requests.exceptions.RequestException):
                error_msg += f"\n\nNetwork error: {str(e)}"
            else:
                error_msg += f"\n\nError: {str(e)}"
            self.show_error(error_msg)
            logging.error(f"Voiceover generation error: {str(e)}")
            return None
    
    def download_images(self, script, output_folder, max_images=5):
        try:
            search_terms = self.extract_keywords(script)
            if not search_terms:
                search_terms = [self.content_type_combo.get()]
            enhanced_terms = [f"{term} high quality" for term in search_terms[:3]]
            fallback_created = False
            for term in enhanced_terms:
                if not self.running:
                    break
                try:
                    downloaded = self.download_google_images(term, min(2, max_images), output_folder)
                    if downloaded == 0 and not fallback_created:
                        self.create_fallback_image(term, output_folder)
                        fallback_created = True
                        max_images -= 1
                except Exception as e:
                    logging.warning(f"Failed to download images for {term}: {str(e)}")
                    if not fallback_created and max_images > 0:
                        self.create_fallback_image(term, output_folder)
                        fallback_created = True
                        max_images -= 1
            image_files = [f for f in os.listdir(output_folder) if f.lower().endswith(('.png', '.jpg', '.jpeg'))]
            if not image_files:
                raise ValueError("Failed to download any images for the video")
        except Exception as e:
            logging.error(f"Image download failed: {str(e)}")
            raise

    def download_google_images(self, query, count, output_dir):
        downloaded = 0
        try:
            params = {
                "q": query,
                "tbm": "isch",
                "hl": "en",
                "tbs": "isz:l"
            }
            response = requests.get(
                "https://www.google.com/search",
                params=params,
                headers={"User-Agent": "Mozilla/5.0"}
            )
            response.raise_for_status()
            soup = BeautifulSoup(response.text, 'html.parser')
            images = soup.find_all('img', limit=count+1)
            for img in images[1:]:
                if downloaded >= count:
                    break
                try:
                    img_url = img['src']
                    if img_url.startswith('http'):
                        img_data = requests.get(img_url, timeout=15).content
                        with Image.open(io.BytesIO(img_data)) as img_test:
                            img_test.verify()
                        img_path = os.path.join(output_dir, f"{query}_{downloaded}.jpg")
                        with open(img_path, 'wb') as f:
                            f.write(img_data)
                        downloaded += 1
                except Exception as e:
                    logging.warning(f"Error downloading image: {str(e)}")
                    continue
        except Exception as e:
            logging.warning(f"Google image download failed for {query}: {str(e)}")
            return downloaded
        return downloaded
    
    def create_fallback_image(self, term, output_dir):
        try:
            img = Image.new('RGB', (1920, 1080), color=(30, 30, 40))
            d = ImageDraw.Draw(img)
            try:
                font_large = ImageFont.truetype("arial.ttf", 72)
                font_small = ImageFont.truetype("arial.ttf", 36)
            except:
                font_large = ImageFont.load_default()
                font_small = ImageFont.load_default()
            for y in range(1080):
                r = int(30 + (y/1080)*20)
                g = int(30 + (y/1080)*20)
                b = int(40 + (y/1080)*20)
                d.line([(0, y), (1920, y)], fill=(r, g, b))
            text = term.replace("high quality", "").strip()
            w, h = d.textsize(text, font=font_large)
            d.text(((1920-w)/2, (1080-h)/2), text, fill=(255, 255, 255), font=font_large)
            d.text((50, 1000), "Generated by AI Content Creator Pro", fill=(200, 200, 200), font=font_small)
            fallback_path = os.path.join(output_dir, f"fallback_{term}.jpg")
            img.save(fallback_path)
            logging.info(f"Created professional fallback image at {fallback_path}")
        except Exception as e:
            logging.error(f"Failed to create fallback image: {str(e)}")
            raise
    
    def create_video_with_effects(self, image_folder, audio_file):
        try:
            image_files = self.get_valid_images(image_folder)
            if not image_files:
                raise ValueError("No valid images available for video creation")
            output_path = self.get_output_path()
            try:
                audio_clip = AudioFileClip(audio_file)
                total_duration = audio_clip.duration
            except Exception as e:
                raise ValueError(f"Invalid audio file: {str(e)}")
            duration_per_image = max(self.img_duration_var.get(), total_duration / len(image_files))
            clips = self.create_video_clips(image_files, duration_per_image)
            if not clips:
                raise ValueError("Could not create any valid video clips")
            final_video = concatenate_videoclips(clips, method="compose")
            try:
                final_video = final_video.set_audio(audio_clip)
            except AttributeError:
                final_video.audio = audio_clip
            self.write_video_file(final_video, output_path)
            return output_path
        except Exception as e:
            raise Exception(f"Video creation failed: {str(e)}")

    def get_valid_images(self, image_folder):
        valid_images = []
        for f in os.listdir(image_folder):
            if f.lower().endswith(('.png', '.jpg', '.jpeg')):
                img_path = os.path.join(image_folder, f)
                try:
                    with Image.open(img_path) as img:
                        img.verify()
                    try:
                        ImageSequenceClip([img_path], durations=[1])
                        valid_images.append(img_path)
                    except:
                        continue
                except:
                    continue
        return valid_images

    def get_output_path(self):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        content_type = self.content_type_combo.get().replace(" ", "_")
        output_filename = f"AI_Video_{content_type}_{timestamp}.mp4"
        return os.path.join(self.output_var.get(), output_filename)

    def create_video_clips(self, image_files, duration_per_image):
        clips = []
        transition = self.transition_combo.get()
        for i, img in enumerate(sorted(image_files)):
            if not self.running:
                break
            try:
                img_clip = ImageSequenceClip([img], durations=[duration_per_image])
                img_clip = img_clip.resize(height=1080)
                zoom_duration = duration_per_image * 0.8
                zoom_factor = self.zoom_var.get()
                zoomed_clip = img_clip.fl_time(lambda t: min(t, zoom_duration))
                zoomed_clip = zoomed_clip.fx(vfx.zoom_in, 
                                           factor=zoom_factor,
                                           zoom_center=(0.5, 0.5))
                if i > 0 and transition != "None":
                    if transition == "Crossfade":
                        zoomed_clip = zoomed_clip.crossfadein(0.5)
                    elif transition == "Fade to Black":
                        if clips:
                            clips[-1] = clips[-1].fx(vfx.fadeout, 0.5)
                        zoomed_clip = zoomed_clip.fx(vfx.fadein, 0.5)
                clips.append(zoomed_clip)
            except Exception as e:
                logging.warning(f"Error processing {img}, using simple clip: {str(e)}")
                try:
                    simple_clip = ImageSequenceClip([img], durations=[duration_per_image])
                    simple_clip = simple_clip.resize(height=1080)
                    clips.append(simple_clip)
                except:
                    logging.warning(f"Failed to create clip for {img}, skipping")
                    continue
        return clips

    def write_video_file(self, final_video, output_path):
        try:
            final_video.write_videofile(
                output_path,
                fps=24,
                codec="libx264",
                audio_codec="aac",
                threads=4,
                bitrate="8000k",
                preset='slow',
                ffmpeg_params=[
                    '-crf', '18',
                    '-pix_fmt', 'yuv420p',
                    '-movflags', '+faststart'
                ]
            )
        except Exception as e:
            logging.warning(f"High quality render failed, trying faster settings: {str(e)}")
            final_video.write_videofile(
                output_path,
                fps=24,
                codec="libx264",
                audio_codec="aac",
                threads=2,
                preset='fast',
                ffmpeg_params=['-crf', '23']
            )

    def cleanup_temp_files(self, image_folder, audio_path):
        try:
            if os.path.exists(audio_path):
                os.remove(audio_path)
            script_path = os.path.join(self.output_dir, "generated_script.docx")
            if os.path.exists(script_path):
                os.remove(script_path)
            if os.path.exists(image_folder):
                for file in os.listdir(image_folder):
                    file_path = os.path.join(image_folder, file)
                    try:
                        if os.path.isfile(file_path):
                            os.remove(file_path)
                    except Exception as e:
                        logging.warning(f"Could not delete {file_path}: {str(e)}")
                try:
                    os.rmdir(image_folder)
                except Exception as e:
                    logging.warning(f"Could not remove directory {image_folder}: {str(e)}")
        except Exception as e:
            logging.warning(f"Cleanup failed: {str(e)}")
    
    def extract_keywords(self, text):
        words = re.findall(r'\b\w{4,}\b', text.lower())
        filtered_words = [w for w in words if w not in [
            'this', 'that', 'with', 'have', 'they', 'what', 'your', 'will',
            'when', 'like', 'just', 'about', 'some', 'from', 'were', 'them'
        ]]
        return list(set(filtered_words))[:10]

if __name__ == "__main__":
    root = tk.Tk()
    app = VideoCreatorApp(root)
    root.mainloop()